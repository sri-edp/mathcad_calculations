from flask import Blueprint, render_template, redirect, url_for, flash, request, jsonify, send_file
from flask_login import current_user, login_required
from app import db
from app.models.calculation_sheet import CalculationSheet, CalculationRevision
from app.models.template import Template, CustomFunction
from datetime import datetime
import json
import os
from werkzeug.utils import secure_filename
import numpy as np
import sympy as sp
from reportlab.pdfgen import canvas
from docx import Document
import matplotlib

matplotlib.use('Agg')  # Use non-interactive backend
import matplotlib.pyplot as plt
import io
import base64
from app.utils.math_engine import evaluate_expression, solve_equation, generate_plot, differentiate, integrate
from app.utils.unit_converter import convert_unit, parse_quantity, evaluate_with_units
from config import Config

calc_bp = Blueprint('calculations', __name__, url_prefix='/calculations')


@calc_bp.route('/new', methods=['GET', 'POST'])
@login_required
def new_calculation():
    """Create a new calculation sheet"""
    if request.method == 'POST':
        try:
            data = request.json

            # Validate required fields
            title = data.get('title', 'Untitled Calculation').strip()
            if not title:
                title = 'Untitled Calculation'

            new_sheet = CalculationSheet(
                title=title,
                description=data.get('description', ''),
                user_id=current_user.id,
                is_template=False,
                is_public=data.get('is_public', False)
            )

            new_sheet.set_content(data.get('content', {'blocks': []}))
            new_sheet.set_variables(data.get('variables', {}))

            db.session.add(new_sheet)
            db.session.commit()

            return jsonify({
                'status': 'success',
                'message': 'Calculation sheet created successfully',
                'id': new_sheet.id
            }), 201

        except Exception as e:
            db.session.rollback()
            return jsonify({
                'status': 'error',
                'message': f'Failed to create calculation: {str(e)}'
            }), 400

    # Handle GET request - show the calculation editor page
    template_id = request.args.get('template_id')
    template = None

    if template_id:
        template = Template.query.get_or_404(template_id)
        # Check if user has access to this template
        if not template.is_published and template.user_id != current_user.id:
            flash('You do not have access to this template', 'error')
            return redirect(url_for('calculations.new_calculation'))

    return render_template(
        'calculations/editor.html',
        title='New Calculation',
        template=template,
        sheet=None
    )


@calc_bp.route('/<int:id>', methods=['GET'])
@login_required
def view_calculation(id):
    """View a calculation sheet"""
    sheet = CalculationSheet.query.get_or_404(id)

    # Check if user has permission to view
    if sheet.user_id != current_user.id and not sheet.is_public:
        flash('You do not have permission to view this calculation', 'error')
        return redirect(url_for('dashboard.index'))

    return render_template(
        'calculations/view.html',
        title=sheet.title,
        sheet=sheet
    )


@calc_bp.route('/<int:id>/edit', methods=['GET', 'PUT'])
@login_required
def edit_calculation(id):
    """Edit a calculation sheet"""
    sheet = CalculationSheet.query.get_or_404(id)

    # Check if user has permission to edit
    if sheet.user_id != current_user.id:
        flash('You do not have permission to edit this calculation', 'error')
        return redirect(url_for('calculations.view_calculation', id=id))

    if request.method == 'PUT':
        try:
            data = request.json

            # Create a revision record before updating
            revision = CalculationRevision(
                sheet_id=sheet.id,
                user_id=current_user.id,
                version=sheet.version
            )

            changes = {
                'previous_content': sheet.content,
                'previous_variables': sheet.variables,
                'previous_title': sheet.title,
                'previous_description': sheet.description,
                'timestamp': datetime.utcnow().isoformat()
            }
            revision.set_changes(changes)

            # Update the sheet
            sheet.title = data.get('title', sheet.title).strip() or 'Untitled Calculation'
            sheet.description = data.get('description', sheet.description)
            sheet.is_public = data.get('is_public', sheet.is_public)
            sheet.version += 1

            # Validate and set content
            content = data.get('content', sheet.get_content())
            if not isinstance(content, dict):
                content = {'blocks': []}
            sheet.set_content(content)

            # Validate and set variables
            variables = data.get('variables', sheet.get_variables())
            if not isinstance(variables, dict):
                variables = {}
            sheet.set_variables(variables)

            # Save changes
            db.session.add(revision)
            db.session.commit()

            return jsonify({
                'status': 'success',
                'message': 'Calculation sheet updated successfully',
                'version': sheet.version
            })

        except Exception as e:
            db.session.rollback()
            return jsonify({
                'status': 'error',
                'message': f'Failed to update calculation: {str(e)}'
            }), 400

    # Handle GET request - show the calculation editor page
    return render_template(
        'calculations/editor.html',
        title=f'Edit: {sheet.title}',
        sheet=sheet
    )


@calc_bp.route('/<int:id>/delete', methods=['DELETE'])
@login_required
def delete_calculation(id):
    """Delete a calculation sheet"""
    sheet = CalculationSheet.query.get_or_404(id)

    # Check if user has permission to delete
    if sheet.user_id != current_user.id:
        return jsonify({
            'status': 'error',
            'message': 'You do not have permission to delete this calculation'
        }), 403

    try:
        # Delete associated revisions first
        CalculationRevision.query.filter_by(sheet_id=sheet.id).delete()

        # Delete the sheet
        db.session.delete(sheet)
        db.session.commit()

        return jsonify({
            'status': 'success',
            'message': 'Calculation sheet deleted successfully'
        })

    except Exception as e:
        db.session.rollback()
        return jsonify({
            'status': 'error',
            'message': f'Failed to delete calculation: {str(e)}'
        }), 500


@calc_bp.route('/<int:id>/export/<format>', methods=['GET'])
@login_required
def export_calculation(id, format):
    """Export a calculation sheet as PDF or Word"""
    sheet = CalculationSheet.query.get_or_404(id)

    # Check if user has permission to view
    if sheet.user_id != current_user.id and not sheet.is_public:
        flash('You do not have permission to export this calculation', 'error')
        return redirect(url_for('dashboard.index'))

    try:
        if format == 'pdf':
            return _export_as_pdf(sheet)
        elif format == 'word':
            return _export_as_word(sheet)
        else:
            flash('Unsupported export format', 'error')
            return redirect(url_for('calculations.view_calculation', id=id))

    except Exception as e:
        flash(f'Export failed: {str(e)}', 'error')
        return redirect(url_for('calculations.view_calculation', id=id))


def _export_as_pdf(sheet):
    """Generate PDF export of calculation sheet"""
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch

    # Create filename and ensure directory exists
    filename = f"{secure_filename(sheet.title)}.pdf"
    filepath = os.path.join(Config.PDF_EXPORT_FOLDER, filename)
    os.makedirs(os.path.dirname(filepath), exist_ok=True)

    # Create PDF document
    doc = SimpleDocTemplate(filepath, pagesize=A4)
    styles = getSampleStyleSheet()
    story = []

    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Title'],
        fontSize=24,
        spaceAfter=30,
    )

    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=16,
        spaceAfter=12,
    )

    # Title
    story.append(Paragraph(sheet.title, title_style))
    story.append(Spacer(1, 20))

    # Metadata
    if sheet.description:
        story.append(Paragraph("Description", heading_style))
        story.append(Paragraph(sheet.description, styles['Normal']))
        story.append(Spacer(1, 12))

    # Author and dates
    metadata_data = [
        ['Author:', f"{sheet.author.first_name} {sheet.author.last_name}"],
        ['Created:', sheet.created_at.strftime('%B %d, %Y at %H:%M')],
        ['Last Updated:', sheet.updated_at.strftime('%B %d, %Y at %H:%M')],
        ['Version:', str(sheet.version)]
    ]

    metadata_table = Table(metadata_data, colWidths=[1.5 * inch, 4 * inch])
    metadata_table.setStyle(TableStyle([
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
    ]))

    story.append(metadata_table)
    story.append(Spacer(1, 20))

    # Variables section
    variables = sheet.get_variables()
    if variables:
        story.append(Paragraph("Variables", heading_style))

        var_data = [['Variable', 'Value', 'Unit', 'Description']]
        for name, var in variables.items():
            var_data.append([
                name,
                str(var.get('value', '')),
                var.get('unit', ''),
                var.get('description', '')
            ])

        var_table = Table(var_data, colWidths=[1 * inch, 1.5 * inch, 1 * inch, 3 * inch])
        var_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))

        story.append(var_table)
        story.append(Spacer(1, 20))

    # Content blocks
    story.append(Paragraph("Calculation Content", heading_style))

    content = sheet.get_content()
    for i, block in enumerate(content.get('blocks', [])):
        block_type = block.get('type')
        block_data = block.get('data', {})

        if block_type == 'text':
            text_content = block_data.get('text', '')
            if text_content:
                story.append(Paragraph(text_content, styles['Normal']))
                story.append(Spacer(1, 12))

        elif block_type == 'equation':
            latex_content = block_data.get('latex', '')
            if latex_content:
                story.append(Paragraph(f"Equation: {latex_content}", styles['Code']))
                story.append(Spacer(1, 12))

        elif block_type == 'calculation':
            input_value = block_data.get('input', '')
            result = block_data.get('result', '')
            if input_value:
                story.append(Paragraph(f"Input: {input_value}", styles['Normal']))
                if result:
                    # Clean HTML tags from result for PDF
                    clean_result = result.replace('<strong>', '').replace('</strong>', '')
                    clean_result = clean_result.replace('<div class="text-danger">', '').replace('</div>', '')
                    story.append(Paragraph(f"Result: {clean_result}", styles['Normal']))
                story.append(Spacer(1, 12))

        elif block_type == 'plot':
            plot_title = block_data.get('title', f'Plot {i + 1}')
            story.append(Paragraph(f"Plot: {plot_title}", styles['Normal']))

            # Handle plot image if available
            image_data = block_data.get('image', '')
            if image_data and image_data.startswith('data:image/png;base64,'):
                try:
                    # Extract base64 data and create image
                    base64_data = image_data.split(',')[1]
                    image_bytes = base64.b64decode(base64_data)

                    # Create temporary image file
                    temp_image_path = os.path.join(Config.PDF_EXPORT_FOLDER, f'temp_plot_{i}.png')
                    with open(temp_image_path, 'wb') as f:
                        f.write(image_bytes)

                    # Add image to PDF
                    img = Image(temp_image_path, width=5 * inch, height=3 * inch)
                    story.append(img)

                    # Clean up temporary file
                    try:
                        os.remove(temp_image_path)
                    except:
                        pass

                except Exception as e:
                    story.append(Paragraph(f"[Plot image could not be rendered: {str(e)}]", styles['Italic']))
            else:
                story.append(Paragraph("[Plot image not available]", styles['Italic']))

            story.append(Spacer(1, 12))

    # Build the PDF
    doc.build(story)

    # Return file for download
    return send_file(filepath, as_attachment=True, download_name=filename)


def _export_as_word(sheet):
    """Generate Word export of calculation sheet"""
    # Create filename and ensure directory exists
    filename = f"{secure_filename(sheet.title)}.docx"
    filepath = os.path.join(Config.WORD_EXPORT_FOLDER, filename)
    os.makedirs(os.path.dirname(filepath), exist_ok=True)

    # Create a new Word document
    doc = Document()

    # Add title
    title = doc.add_heading(sheet.title, 0)
    title.alignment = 1  # Center alignment

    # Add metadata
    doc.add_heading('Document Information', level=1)

    if sheet.description:
        p = doc.add_paragraph()
        p.add_run('Description: ').bold = True
        p.add_run(sheet.description)

    p = doc.add_paragraph()
    p.add_run('Author: ').bold = True
    p.add_run(f"{sheet.author.first_name} {sheet.author.last_name}")

    p = doc.add_paragraph()
    p.add_run('Created: ').bold = True
    p.add_run(sheet.created_at.strftime('%B %d, %Y at %H:%M'))

    p = doc.add_paragraph()
    p.add_run('Last Updated: ').bold = True
    p.add_run(sheet.updated_at.strftime('%B %d, %Y at %H:%M'))

    p = doc.add_paragraph()
    p.add_run('Version: ').bold = True
    p.add_run(str(sheet.version))

    # Add variables section
    variables = sheet.get_variables()
    if variables:
        doc.add_heading('Variables', level=1)

        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'

        # Header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Variable'
        hdr_cells[1].text = 'Value'
        hdr_cells[2].text = 'Unit'
        hdr_cells[3].text = 'Description'

        for name, var in variables.items():
            row_cells = table.add_row().cells
            row_cells[0].text = name
            row_cells[1].text = str(var.get('value', ''))
            row_cells[2].text = var.get('unit', '')
            row_cells[3].text = var.get('description', '')

    # Add content blocks
    doc.add_heading('Calculation Content', level=1)

    content = sheet.get_content()
    for i, block in enumerate(content.get('blocks', [])):
        block_type = block.get('type')
        block_data = block.get('data', {})

        if block_type == 'text':
            text_content = block_data.get('text', '')
            if text_content:
                doc.add_paragraph(text_content)

        elif block_type == 'equation':
            latex_content = block_data.get('latex', '')
            if latex_content:
                p = doc.add_paragraph()
                p.add_run('Equation: ').bold = True
                p.add_run(latex_content).italic = True

        elif block_type == 'calculation':
            input_value = block_data.get('input', '')
            result = block_data.get('result', '')
            if input_value:
                p = doc.add_paragraph()
                p.add_run('Input: ').bold = True
                p.add_run(input_value)

                if result:
                    # Clean HTML tags from result
                    clean_result = result.replace('<strong>', '').replace('</strong>', '')
                    clean_result = clean_result.replace('<div class="text-danger">', '').replace('</div>', '')
                    p = doc.add_paragraph()
                    p.add_run('Result: ').bold = True
                    p.add_run(clean_result)

        elif block_type == 'plot':
            plot_title = block_data.get('title', f'Plot {i + 1}')
            p = doc.add_paragraph()
            p.add_run('Plot: ').bold = True
            p.add_run(plot_title)

            # Handle plot image
            image_data = block_data.get('image', '')
            if image_data and image_data.startswith('data:image/png;base64,'):
                try:
                    # Extract base64 data and create image
                    base64_data = image_data.split(',')[1]
                    image_bytes = base64.b64decode(base64_data)

                    # Create temporary image file
                    temp_image_path = os.path.join(Config.WORD_EXPORT_FOLDER, f'temp_plot_{i}.png')
                    with open(temp_image_path, 'wb') as f:
                        f.write(image_bytes)

                    # Add image to document
                    doc.add_picture(temp_image_path, width=6000000)  # 6 inches

                    # Clean up temporary file
                    try:
                        os.remove(temp_image_path)
                    except:
                        pass

                except Exception as e:
                    doc.add_paragraph(f"[Plot image could not be rendered: {str(e)}]")
            else:
                doc.add_paragraph("[Plot image not available]")

    # Save the Word document
    doc.save(filepath)

    # Return file for download
    return send_file(filepath, as_attachment=True, download_name=filename)


# API Endpoints for mathematical operations
@calc_bp.route('/api/evaluate', methods=['POST'])
@login_required
def api_evaluate():
    """API endpoint to evaluate mathematical expressions with enhanced unit support"""
    try:
        data = request.json
        expression = data.get('expression', '').strip()
        variables = data.get('variables', {})

        if not expression:
            return jsonify({
                'status': 'error',
                'message': 'Expression is required'
            }), 400

        # Enhanced evaluation with unit tracking
        result = evaluate_expression(expression, variables)

        return jsonify({
            'status': 'success',
            'result': result
        })

    except Exception as e:
        return jsonify({
            'status': 'error',
            'message': str(e)
        }), 400


@calc_bp.route('/api/solve', methods=['POST'])
@login_required
def api_solve():
    """API endpoint to solve equations with enhanced unit support"""
    try:
        data = request.json
        equation = data.get('equation', '').strip()
        variables = data.get('variables', {})
        solving_for = data.get('solving_for', '').strip()

        if not equation:
            return jsonify({
                'status': 'error',
                'message': 'Equation is required'
            }), 400

        if not solving_for:
            return jsonify({
                'status': 'error',
                'message': 'Variable to solve for is required'
            }), 400

        # Enhanced equation solving with unit handling
        solutions = solve_equation(equation, solving_for, variables)

        return jsonify({
            'status': 'success',
            'result': solutions
        })

    except Exception as e:
        return jsonify({
            'status': 'error',
            'message': str(e)
        }), 400


@calc_bp.route('/api/differentiate', methods=['POST'])
@login_required
def api_differentiate():
    """API endpoint for symbolic differentiation"""
    try:
        data = request.json
        expression = data.get('expression', '').strip()
        variable = data.get('variable', 'x').strip()
        order = data.get('order', 1)

        if not expression:
            return jsonify({
                'status': 'error',
                'message': 'Expression is required'
            }), 400

        result = differentiate(expression, variable, order)

        return jsonify({
            'status': 'success',
            'result': result
        })

    except Exception as e:
        return jsonify({
            'status': 'error',
            'message': str(e)
        }), 400


@calc_bp.route('/api/integrate', methods=['POST'])
@login_required
def api_integrate():
    """API endpoint for symbolic integration"""
    try:
        data = request.json
        expression = data.get('expression', '').strip()
        variable = data.get('variable', 'x').strip()
        lower_bound = data.get('lower_bound')
        upper_bound = data.get('upper_bound')

        if not expression:
            return jsonify({
                'status': 'error',
                'message': 'Expression is required'
            }), 400

        result = integrate(expression, variable, lower_bound, upper_bound)

        return jsonify({
            'status': 'success',
            'result': result
        })

    except Exception as e:
        return jsonify({
            'status': 'error',
            'message': str(e)
        }), 400


@calc_bp.route('/api/plot', methods=['POST'])
@login_required
def api_plot():
    """API endpoint to generate enhanced plots"""
    try:
        data = request.json
        plot_type = data.get('type', 'line')
        x_data = data.get('x_data', [])
        y_data = data.get('y_data', [])
        title = data.get('title', 'Plot')
        x_label = data.get('x_label', 'X')
        y_label = data.get('y_label', 'Y')

        # Additional plot options
        grid = data.get('grid', True)
        legend = data.get('legend', False)
        color = data.get('color', 'blue')
        style = data.get('style', '-')
        marker = data.get('marker', None)

        if not x_data or not y_data:
            return jsonify({
                'status': 'error',
                'message': 'Both x_data and y_data are required'
            }), 400

        if len(x_data) != len(y_data):
            return jsonify({
                'status': 'error',
                'message': 'x_data and y_data must have the same length'
            }), 400

        # Generate enhanced plot
        image_data = generate_plot(
            plot_type, x_data, y_data, title, x_label, y_label,
            grid=grid, legend=legend, color=color, style=style, marker=marker
        )

        return jsonify({
            'status': 'success',
            'image_data': image_data,
            'plot_type': plot_type,
            'data_points': len(x_data)
        })

    except Exception as e:
        return jsonify({
            'status': 'error',
            'message': str(e)
        }), 400


@calc_bp.route('/api/convert_unit', methods=['POST'])
@login_required
def api_convert_unit():
    """API endpoint for enhanced unit conversion"""
    try:
        data = request.json
        value = data.get('value')
        from_unit = data.get('from_unit', '').strip()
        to_unit = data.get('to_unit', '').strip()

        if value is None:
            return jsonify({
                'status': 'error',
                'message': 'Value is required'
            }), 400

        if not from_unit or not to_unit:
            return jsonify({
                'status': 'error',
                'message': 'Both from_unit and to_unit are required'
            }), 400

        # Handle string input with embedded units
        if isinstance(value, str):
            parsed_value, parsed_unit = parse_quantity(value)
            if parsed_unit and not from_unit:
                from_unit = parsed_unit
            value = parsed_value

        result_value = convert_unit(float(value), from_unit, to_unit)

        return jsonify({
            'status': 'success',
            'result': {
                'value': result_value,
                'unit': to_unit,
                'formatted': f"{result_value} {to_unit}",
                'original_value': value,
                'original_unit': from_unit
            }
        })

    except Exception as e:
        return jsonify({
            'status': 'error',
            'message': str(e)
        }), 400


@calc_bp.route('/<int:id>/copy', methods=['POST'])
@login_required
def copy_calculation(id):
    """Copy a shared calculation to the user's personal collection"""
    source_sheet = CalculationSheet.query.get_or_404(id)

    # Check if the calculation is public or belongs to the user
    if not source_sheet.is_public and source_sheet.user_id != current_user.id:
        return jsonify({
            'status': 'error',
            'message': 'You do not have permission to copy this calculation'
        }), 403

    try:
        # Create a new calculation sheet as a copy
        new_sheet = CalculationSheet(
            title=f"Copy of {source_sheet.title}",
            description=source_sheet.description,
            content=source_sheet.content,
            variables=source_sheet.variables,
            user_id=current_user.id,
            is_template=False,
            is_public=False  # Copies are private by default
        )

        db.session.add(new_sheet)
        db.session.commit()

        return jsonify({
            'status': 'success',
            'message': 'Calculation copied successfully',
            'id': new_sheet.id
        })

    except Exception as e:
        db.session.rollback()
        return jsonify({
            'status': 'error',
            'message': f'Failed to copy calculation: {str(e)}'
        }), 500


@calc_bp.route('/<int:id>/revisions', methods=['GET'])
@login_required
def view_revisions(id):
    """View revision history of a calculation"""
    sheet = CalculationSheet.query.get_or_404(id)

    # Check if user has permission to view
    if sheet.user_id != current_user.id:
        flash('You do not have permission to view this calculation\'s revisions', 'error')
        return redirect(url_for('calculations.view_calculation', id=id))

    # Get all revisions for this sheet
    revisions = CalculationRevision.query.filter_by(
        sheet_id=sheet.id
    ).order_by(CalculationRevision.timestamp.desc()).all()

    return render_template(
        'calculations/revisions.html',
        title=f'Revisions: {sheet.title}',
        sheet=sheet,
        revisions=revisions
    )


@calc_bp.route('/api/validate', methods=['POST'])
@login_required
def api_validate():
    """API endpoint to validate expressions and equations"""
    try:
        data = request.json
        expression = data.get('expression', '').strip()
        expression_type = data.get('type', 'expression')  # 'expression' or 'equation'

        if not expression:
            return jsonify({
                'status': 'error',
                'message': 'Expression is required'
            }), 400

        # Basic validation
        validation_result = {
            'valid': True,
            'errors': [],
            'warnings': [],
            'suggestions': []
        }

        try:
            if expression_type == 'equation':
                if '=' not in expression:
                    validation_result['warnings'].append('Equation should contain an equals sign (=)')
                else:
                    # Try to parse both sides
                    left, right = expression.split('=', 1)
                    sp.sympify(left.strip())
                    sp.sympify(right.strip())
            else:
                # Validate as expression
                sp.sympify(expression)

        except Exception as e:
            validation_result['valid'] = False
            validation_result['errors'].append(f'Syntax error: {str(e)}')

        return jsonify({
            'status': 'success',
            'result': validation_result
        })

    except Exception as e:
        return jsonify({
            'status': 'error',
            'message': str(e)
        }), 400


@calc_bp.errorhandler(404)
def calc_not_found(error):
    return render_template('errors/404.html'), 404


@calc_bp.errorhandler(500)
def calc_internal_error(error):
    db.session.rollback()
    return render_template('errors/500.html'), 500